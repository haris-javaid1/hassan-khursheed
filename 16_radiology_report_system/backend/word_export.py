from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def create_table_format(report):
    """Generate Word document in table format"""
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('RADIOLOGY REPORT', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Add space
    
    # Create patient information table
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    
    # Row 1: UHID and Patient No
    table.rows[0].cells[0].text = 'UHID:'
    table.rows[0].cells[1].text = str(report.uhid or '')
    table.rows[0].cells[2].text = 'Patient No:'
    table.rows[0].cells[3].text = str(report.patient_no or '')
    
    # Row 2: Sl No and Reg No
    table.rows[1].cells[0].text = 'Sl No:'
    table.rows[1].cells[1].text = str(report.sl_no or '')
    table.rows[1].cells[2].text = 'Reg No:'
    table.rows[1].cells[3].text = str(report.reg_no or '')
    
    # Row 3: Patient Name
    table.rows[2].cells[0].text = 'Patient Name:'
    merged_cell = table.rows[2].cells[1].merge(table.rows[2].cells[3])
    merged_cell.text = str(report.patient_name or '')
    
    # Row 4: Age/Sex and Report Date
    table.rows[3].cells[0].text = 'Age/Sex:'
    table.rows[3].cells[1].text = str(report.age_sex or '')
    table.rows[3].cells[2].text = 'Report Date:'
    table.rows[3].cells[3].text = str(report.report_date or '')
    
    # Row 5: Origin and Referred By
    table.rows[4].cells[0].text = 'Origin:'
    table.rows[4].cells[1].text = str(report.origin_ethe or '')
    table.rows[4].cells[2].text = 'Referred By:'
    table.rows[4].cells[3].text = str(report.ref_by or '')
    
    # Row 6: Film No and Scan Time
    table.rows[5].cells[0].text = 'Film No:'
    table.rows[5].cells[1].text = str(report.film_no or '')
    table.rows[5].cells[2].text = 'Scan Time:'
    table.rows[5].cells[3].text = str(report.scan_time or '')
    
    # Row 7: Report Time and TAT
    table.rows[6].cells[0].text = 'Report Time:'
    table.rows[6].cells[1].text = str(report.report_time or '')
    table.rows[6].cells[2].text = 'TAT:'
    table.rows[6].cells[3].text = str(report.tat or '')
    
    doc.add_paragraph()  # Add space
    
    # Add scan type
    scan_heading = doc.add_heading(str(report.scan_type or 'SCAN TYPE'), level=2)
    scan_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Add space
    
    # Add findings
    doc.add_heading('Findings:', level=3)
    doc.add_paragraph(str(report.doctor_description or 'No findings recorded.'))
    
    doc.add_paragraph()  # Add space
    
    # Add impression
    doc.add_heading('Impression:', level=3)
    doc.add_paragraph(str(report.impression or 'No impression recorded.'))
    
    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream


def create_list_format(report):
    """Generate Word document in list format"""
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('RADIOLOGY REPORT', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Patient Information
    doc.add_heading('Patient Information:', level=2)
    doc.add_paragraph(f'• UHID: {report.uhid or "N/A"}')
    doc.add_paragraph(f'• Patient Name: {report.patient_name or "N/A"}')
    doc.add_paragraph(f'• Patient No: {report.patient_no or "N/A"}')
    doc.add_paragraph(f'• Age/Sex: {report.age_sex or "N/A"}')
    doc.add_paragraph(f'• Origin: {report.origin_ethe or "N/A"}')
    doc.add_paragraph(f'• Referred By: {report.ref_by or "N/A"}')
    
    doc.add_paragraph()
    
    # Scan Information
    doc.add_heading('Scan Information:', level=2)
    doc.add_paragraph(f'• Scan Type: {report.scan_type or "N/A"}')
    doc.add_paragraph(f'• Report Date: {report.report_date or "N/A"}')
    doc.add_paragraph(f'• Film No: {report.film_no or "N/A"}')
    doc.add_paragraph(f'• Scan Time: {report.scan_time or "N/A"}')
    doc.add_paragraph(f'• Report Time: {report.report_time or "N/A"}')
    doc.add_paragraph(f'• TAT: {report.tat or "N/A"}')
    
    doc.add_paragraph()
    
    # Findings
    doc.add_heading('Findings:', level=2)
    doc.add_paragraph(str(report.doctor_description or 'No findings recorded.'))
    
    doc.add_paragraph()
    
    # Impression
    doc.add_heading('Impression:', level=2)
    doc.add_paragraph(str(report.impression or 'No impression recorded.'))
    
    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream


def create_detailed_format(report):
    """Generate Word document in detailed format"""
    
    doc = Document()
    
    # Main title
    title = doc.add_heading('DETAILED RADIOLOGY REPORT', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Administrative Information Section
    doc.add_heading('Administrative Information', level=2)
    
    admin_table = doc.add_table(rows=4, cols=2)
    admin_table.style = 'Light Grid Accent 1'
    
    admin_table.rows[0].cells[0].text = 'UHID'
    admin_table.rows[0].cells[1].text = str(report.uhid or '')
    admin_table.rows[1].cells[0].text = 'Sl No.'
    admin_table.rows[1].cells[1].text = str(report.sl_no or '')
    admin_table.rows[2].cells[0].text = 'Registration No.'
    admin_table.rows[2].cells[1].text = str(report.reg_no or '')
    admin_table.rows[3].cells[0].text = 'Patient No.'
    admin_table.rows[3].cells[1].text = str(report.patient_no or '')
    
    doc.add_paragraph()
    
    # Patient Demographics Section
    doc.add_heading('Patient Demographics', level=2)
    
    patient_table = doc.add_table(rows=4, cols=2)
    patient_table.style = 'Light Grid Accent 1'
    
    patient_table.rows[0].cells[0].text = 'Patient Name'
    patient_table.rows[0].cells[1].text = str(report.patient_name or '')
    patient_table.rows[1].cells[0].text = 'Age/Sex'
    patient_table.rows[1].cells[1].text = str(report.age_sex or '')
    patient_table.rows[2].cells[0].text = 'Origin'
    patient_table.rows[2].cells[1].text = str(report.origin_ethe or '')
    patient_table.rows[3].cells[0].text = 'Referred By'
    patient_table.rows[3].cells[1].text = str(report.ref_by or '')
    
    doc.add_paragraph()
    
    # Examination Details Section
    doc.add_heading('Examination Details', level=2)
    
    exam_table = doc.add_table(rows=6, cols=2)
    exam_table.style = 'Light Grid Accent 1'
    
    exam_table.rows[0].cells[0].text = 'Scan Type'
    exam_table.rows[0].cells[1].text = str(report.scan_type or '')
    exam_table.rows[1].cells[0].text = 'Report Date'
    exam_table.rows[1].cells[1].text = str(report.report_date or '')
    exam_table.rows[2].cells[0].text = 'Film No.'
    exam_table.rows[2].cells[1].text = str(report.film_no or '')
    exam_table.rows[3].cells[0].text = 'Scan Time'
    exam_table.rows[3].cells[1].text = str(report.scan_time or '')
    exam_table.rows[4].cells[0].text = 'Report Time'
    exam_table.rows[4].cells[1].text = str(report.report_time or '')
    exam_table.rows[5].cells[0].text = 'TAT'
    exam_table.rows[5].cells[1].text = str(report.tat or '')
    
    doc.add_paragraph()
    
    # Clinical Findings
    doc.add_heading('Clinical Findings', level=2)
    doc.add_paragraph(str(report.doctor_description or 'No findings recorded.'))
    
    doc.add_paragraph()
    
    # Diagnostic Impression
    doc.add_heading('Diagnostic Impression', level=2)
    doc.add_paragraph(str(report.impression or 'No impression recorded.'))
    
    doc.add_paragraph()
    
    # Footer
    doc.add_paragraph('_' * 80)
    footer = doc.add_paragraph(f'Report Generated: {report.created_at}')
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream