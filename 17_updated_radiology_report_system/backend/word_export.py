from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def create_table_format(report):
    """Generate Word document in table format"""
    doc = Document()
    
    title = doc.add_heading('RADIOLOGY REPORT', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    
    table.rows[0].cells[0].text = 'UHID:'
    table.rows[0].cells[1].text = str(report.uhid or '')
    table.rows[0].cells[2].text = 'Patient No:'
    table.rows[0].cells[3].text = str(report.patient_no or '')
    
    table.rows[1].cells[0].text = 'Sl No:'
    table.rows[1].cells[1].text = str(report.sl_no or '')
    table.rows[1].cells[2].text = 'Reg No:'
    table.rows[1].cells[3].text = str(report.reg_no or '')
    
    table.rows[2].cells[0].text = 'Patient Name:'
    merged_cell = table.rows[2].cells[1].merge(table.rows[2].cells[3])
    merged_cell.text = str(report.patient_name or '')
    
    table.rows[3].cells[0].text = 'Age/Sex:'
    table.rows[3].cells[1].text = str(report.age_sex or '')
    table.rows[3].cells[2].text = 'Report Date:'
    table.rows[3].cells[3].text = str(report.report_date or '')
    
    table.rows[4].cells[0].text = 'Origin:'
    table.rows[4].cells[1].text = str(report.origin_ethe or '')
    table.rows[4].cells[2].text = 'Referred By:'
    table.rows[4].cells[3].text = str(report.ref_by or '')
    
    table.rows[5].cells[0].text = 'Film No:'
    table.rows[5].cells[1].text = str(report.film_no or '')
    table.rows[5].cells[2].text = 'Scan Time:'
    table.rows[5].cells[3].text = str(report.scan_time or '')
    
    table.rows[6].cells[0].text = 'Report Time:'
    table.rows[6].cells[1].text = str(report.report_time or '')
    table.rows[6].cells[2].text = 'TAT:'
    table.rows[6].cells[3].text = str(report.tat or '')
    
    doc.add_paragraph()
    
    scan_heading = doc.add_heading(str(report.scan_type or 'SCAN TYPE'), level=2)
    scan_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    doc.add_heading('Findings:', level=3)
    doc.add_paragraph(str(report.doctor_description or 'No findings recorded.'))
    doc.add_paragraph()
    
    doc.add_heading('Impression:', level=3)
    doc.add_paragraph(str(report.impression or 'No impression recorded.'))
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream