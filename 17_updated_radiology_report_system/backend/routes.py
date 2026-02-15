from fastapi import APIRouter, Depends, HTTPException, Request, File, UploadFile
from fastapi.responses import StreamingResponse, Response, HTMLResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import Optional
from database import get_db
from models import RadiologyReport
from word_export import create_table_format
from qr_generator import generate_serial_number, create_qr_code
from docx import Document
from io import BytesIO
from datetime import datetime

templates = Jinja2Templates(directory="templates")
router = APIRouter()

class ReportCreate(BaseModel):
    uhid: Optional[str] = None
    sl_no: Optional[str] = None
    reg_no: Optional[str] = None
    patient_no: Optional[str] = None
    patient_name: str
    report_date: Optional[str] = None
    age_sex: Optional[str] = None
    origin_ethe: Optional[str] = None
    ref_by: Optional[str] = None
    film_no: Optional[str] = None
    scan_time: Optional[str] = None
    report_time: Optional[str] = None
    tat: Optional[str] = None
    scan_type: Optional[str] = None
    doctor_description: Optional[str] = None
    impression: Optional[str] = None


# Create a new radiology report with auto-generated serial number
@router.post("/reports")
def create_report(report: ReportCreate, db: Session = Depends(get_db)):
    try:
        serial_number = generate_serial_number()
        
        new_report = RadiologyReport(
            **report.dict(),
            serial_number=serial_number
        )
        
        db.add(new_report)
        db.commit()
        db.refresh(new_report)
        
        return {
            "success": True,
            "id": new_report.id,
            "serial_number": serial_number,
            "message": "Report created successfully"
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


# Get all reports for displaying in reports table
@router.get("/reports")
def get_all_reports(db: Session = Depends(get_db)):
    try:
        reports = db.query(RadiologyReport).order_by(
            RadiologyReport.created_at.desc()
        ).all()
        
        reports_list = []
        for report in reports:
            reports_list.append({
                "id": report.id,
                "serial_number": report.serial_number,
                "uhid": report.uhid,
                "patient_name": report.patient_name,
                "patient_no": report.patient_no,
                "report_date": str(report.report_date) if report.report_date else None,
                "age_sex": report.age_sex,
                "scan_type": report.scan_type,
                "created_at": str(report.created_at)
            })
        
        return reports_list
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# Get single report by serial number for editing
@router.get("/reports/serial/{serial_number}")
def get_report_by_serial(serial_number: str, db: Session = Depends(get_db)):
    try:
        report = db.query(RadiologyReport).filter(
            RadiologyReport.serial_number == serial_number
        ).first()
        
        if not report:
            raise HTTPException(status_code=404, detail="Report not found")
        
        return {
            "id": report.id,
            "serial_number": report.serial_number,
            "uhid": report.uhid,
            "sl_no": report.sl_no,
            "reg_no": report.reg_no,
            "patient_no": report.patient_no,
            "patient_name": report.patient_name,
            "report_date": str(report.report_date) if report.report_date else None,
            "age_sex": report.age_sex,
            "origin_ethe": report.origin_ethe,
            "ref_by": report.ref_by,
            "film_no": report.film_no,
            "scan_time": report.scan_time,
            "report_time": report.report_time,
            "tat": report.tat,
            "scan_type": report.scan_type,
            "doctor_description": report.doctor_description,
            "impression": report.impression
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# Update existing report with new data
@router.put("/reports/serial/{serial_number}")
def update_report(serial_number: str, report: ReportCreate, db: Session = Depends(get_db)):
    try:
        db_report = db.query(RadiologyReport).filter(
            RadiologyReport.serial_number == serial_number
        ).first()
        
        if not db_report:
            raise HTTPException(status_code=404, detail="Report not found")
        
        for key, value in report.dict().items():
            if value is not None:
                setattr(db_report, key, value)
        
        db.commit()
        return {"success": True, "message": "Report updated successfully"}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


# Delete a report from database
@router.delete("/reports/serial/{serial_number}")
def delete_report(serial_number: str, db: Session = Depends(get_db)):
    try:
        report = db.query(RadiologyReport).filter(
            RadiologyReport.serial_number == serial_number
        ).first()
        
        if not report:
            raise HTTPException(status_code=404, detail="Report not found")
        
        db.delete(report)
        db.commit()
        return {"success": True, "message": "Report deleted successfully"}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


# Upload Word document and extract data to update report in database
@router.post("/reports/serial/{serial_number}/upload")
async def upload_file(serial_number: str, file: UploadFile = File(...), db: Session = Depends(get_db)):
    try:
        if not file.filename.endswith(".docx"):
            raise HTTPException(status_code=400, detail="Only .docx files allowed")

        content = await file.read()
        document = Document(BytesIO(content))
        extracted = {}

        # Extract table data
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    extracted[cells[0].replace(":", "").strip()] = cells[1].strip()
                if len(cells) >= 4:
                    extracted[cells[2].replace(":", "").strip()] = cells[3].strip()
            break

        # Extract scan type and sections
        scan_type = None
        findings_text = ""
        impression_text = ""
        current_section = None
        last_text = ""

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if text.startswith("Findings:"):
                scan_type = last_text
                current_section = "Findings"
                continue
            elif text.startswith("Impression:"):
                current_section = "Impression"
                continue

            if current_section == "Findings":
                findings_text += text + "\n"
            elif current_section == "Impression":
                impression_text += text + "\n"

            last_text = text

        # Parse report date
        report_date = None
        if extracted.get("Report Date"):
            for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
                try:
                    report_date = datetime.strptime(extracted["Report Date"], fmt).date()
                    break
                except:
                    continue

        # Prepare data for database
        report_data = {
            "uhid": extracted.get("UHID"),
            "sl_no": extracted.get("Sl No"),
            "reg_no": extracted.get("Reg No"),
            "patient_no": extracted.get("Patient No"),
            "patient_name": extracted.get("Patient Name", "Uploaded Patient"),
            "age_sex": extracted.get("Age/Sex"),
            "origin_ethe": extracted.get("Origin"),
            "ref_by": extracted.get("Referred By"),
            "film_no": extracted.get("Film No"),
            "scan_time": extracted.get("Scan Time"),
            "report_time": extracted.get("Report Time"),
            "tat": extracted.get("TAT"),
            "scan_type": scan_type,
            "report_date": report_date,
            "doctor_description": findings_text.strip(),
            "impression": impression_text.strip(),
        }

        # Update existing report
        report = db.query(RadiologyReport).filter(
            RadiologyReport.serial_number == serial_number
        ).first()
        
        if report:
            for k, v in report_data.items():
                if v is not None:
                    setattr(report, k, v)
            db.commit()
            return {"success": True, "message": "Report updated from file"}

        # Create new report if not exists
        new_report = RadiologyReport(
            serial_number=serial_number,
            **report_data
        )
        db.add(new_report)
        db.commit()
        return {"success": True, "message": "New report created from file"}

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


# Generate QR code on-the-fly for report viewing
@router.get("/reports/serial/{serial_number}/qr")
def get_qr(serial_number: str, db: Session = Depends(get_db)):
    try:
        report = db.query(RadiologyReport).filter(
            RadiologyReport.serial_number == serial_number
        ).first()
        
        if not report:
            raise HTTPException(status_code=404, detail="Report not found")
        
        # Generate QR code dynamically (not stored in database)
        qr_bytes = create_qr_code(serial_number)
        
        return Response(content=qr_bytes, media_type="image/png")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# Export report as Word document
@router.get("/reports/{report_id}/export")
def export_report(report_id: int, db: Session = Depends(get_db)):
    try:
        report = db.query(RadiologyReport).filter(
            RadiologyReport.id == report_id
        ).first()
        
        if not report:
            raise HTTPException(status_code=404, detail="Report not found")
        
        file_stream = create_table_format(report)
        filename = f"radiology_report_{report_id}.docx"
        
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# Display report online when QR code is scanned
@router.get("/view/{serial_number}", response_class=HTMLResponse)
def view_online(serial_number: str, request: Request, db: Session = Depends(get_db)):
    try:
        report = db.query(RadiologyReport).filter(
            RadiologyReport.serial_number == serial_number
        ).first()

        if not report:
            return HTMLResponse(content="<h1>Report not found</h1>", status_code=404)

        return templates.TemplateResponse("report_view.html", {
            "request": request,
            "report": report
        })

    except Exception as e:
        return HTMLResponse(content=f"<h1>Error: {str(e)}</h1>", status_code=500)